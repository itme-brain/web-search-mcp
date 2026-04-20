import asyncio
import fnmatch
import hashlib
import io
import json
import logging
import mimetypes
import os
import re
import time
from inspect import isawaitable
from collections import defaultdict, deque
from urllib.parse import parse_qs, urlencode, urljoin, urlparse, urlunparse

import httpx
import pysbd
import trafilatura
from fastmcp import Context, FastMCP
from pydantic_settings import BaseSettings
from url_normalize import url_normalize
from flashrank import Ranker, RerankRequest
from docx import Document as DocxDocument
from pypdf import PdfReader
from starlette.requests import Request
from starlette.responses import JSONResponse

logging.basicConfig(
    level=os.environ.get("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
)
log = logging.getLogger("web-search-mcp")

mcp = FastMCP("Web Search", version="0.0.1")


class Settings(BaseSettings):
    searxng_url: str = "http://searxng:8080"
    crawl4ai_url: str = "http://crawl4ai:11235"
    rerank_model: str = "ms-marco-MiniLM-L-12-v2"
    request_timeout: int = 30
    max_results: int = 10
    max_scrape: int = 5


settings = Settings()

SEARXNG_URL = settings.searxng_url
CRAWL4AI_URL = settings.crawl4ai_url
RERANK_MODEL = settings.rerank_model
REQUEST_TIMEOUT = settings.request_timeout
MAX_RESULTS = settings.max_results
MAX_SCRAPE = settings.max_scrape

# Internal constants — not user-configurable.
_RERANK_MAX_LENGTH = 512
_HTTP_TIMEOUT = max(REQUEST_TIMEOUT // 2, 10)
_MAX_CONTENT_CHARS = 8000
_DEDUP_SIMILARITY = 0.75
_TOP_CHUNKS = 3
_MAX_CHUNKS_PER_PAGE = 10
_MAX_EXTRACT_URLS = 20
_MAX_MAP_URLS = 50
_MAX_MAP_DEPTH = 2

STATE_SCRAPE_CACHE = "scrape_cache"
STATE_QUERY_CACHE = "query_cache"
STATE_SEEN_URLS = "seen_urls"
STATE_EXTRACT_CACHE = "extract_cache"

VALID_TIME_RANGES = frozenset({"day", "week", "month", "year"})
VALID_MODES = frozenset({"balanced", "deep"})
VALID_CATEGORIES = frozenset({
    "general", "news", "images", "videos", "music",
    "files", "science", "social media", "it",
})

_SENTENCE_SPLITTER = pysbd.Segmenter(language="en", clean=False)
_TRACKING_PARAMS = frozenset({
    "utm_source", "utm_medium", "utm_campaign", "utm_term", "utm_content",
    "ref", "fbclid", "gclid", "dclid", "msclkid", "mc_cid", "mc_eid",
})
_WORD_SPLIT = re.compile(r"\W+")
_WHITESPACE = re.compile(r"\s+")


def _warning(error_type: str, source: str, detail: str) -> dict:
    """Create a structured warning dict for programmatic error handling."""
    return {"type": error_type, "source": source, "detail": detail}


def _normalize_url(url: str) -> str:
    normalized = url_normalize(url)
    parsed = urlparse(normalized)
    host = parsed.hostname or ""
    if host.startswith("www."):
        host = host[4:]
    params = {k: v for k, v in parse_qs(parsed.query).items() if k not in _TRACKING_PARAMS}
    return urlunparse((parsed.scheme, host, parsed.path.rstrip("/"), "", urlencode(params, doseq=True), ""))


def _dedup_results(results: list[dict]) -> list[dict]:
    seen_urls: set[str] = set()
    seen_titles: set[tuple[str, str]] = set()
    deduped: list[dict] = []
    for r in results:
        norm = _normalize_url(r.get("url", ""))
        domain = _domain_from_url(r.get("url", "")).lower()
        title = _normalize_title(r.get("title", ""))
        title_key = (domain, title)
        if norm in seen_urls:
            continue
        if title and title_key in seen_titles:
            continue
        seen_urls.add(norm)
        if title:
            seen_titles.add(title_key)
        deduped.append(r)
    return deduped


def _normalize_title(title: str) -> str:
    normalized = _WHITESPACE.sub(" ", title.strip().lower())
    normalized = re.sub(r"[^a-z0-9 ]+", "", normalized)
    normalized = _WHITESPACE.sub(" ", normalized)
    return normalized.strip()


def _word_set(text: str) -> set[str]:
    return set(_WORD_SPLIT.split(text.lower()))


def _dedup_chunks(chunks: list[str], entry_map: list[int]) -> tuple[list[str], list[int]]:
    kept_chunks: list[str] = []
    kept_entries: list[int] = []
    seen_words: list[set[str]] = []
    for chunk, eidx in zip(chunks, entry_map):
        words = _word_set(chunk)
        if not words:
            continue
        is_dup = False
        for prev in seen_words:
            intersection = len(words & prev)
            union = len(words | prev)
            if union and intersection / union >= _DEDUP_SIMILARITY:
                is_dup = True
                break
        if not is_dup:
            kept_chunks.append(chunk)
            kept_entries.append(eidx)
            seen_words.append(words)
    return kept_chunks, kept_entries


def _query_cache_key(query: str, num_results: int, scrape_top: int, time_range: str | None) -> str:
    raw = json.dumps([query.lower().strip(), num_results, scrape_top, time_range], sort_keys=True)
    return hashlib.sha256(raw.encode()).hexdigest()


def _response_cache_key(query: str, num_results: int, scrape_top: int, time_range: str | None) -> str:
    return _query_cache_key(query, num_results, scrape_top, time_range)


def _normalize_time_range(time_range: str | None) -> str | None:
    if time_range in (None, "null", "none", "None", ""):
        return None
    normalized = time_range.strip().strip('"').lower()
    if normalized not in VALID_TIME_RANGES:
        raise ValueError(f"invalid time_range: {time_range!r}. Expected one of {sorted(VALID_TIME_RANGES)}")
    return normalized


def _validate_query(query: str) -> str:
    normalized = query.strip()
    if not normalized:
        raise ValueError("query must not be empty")
    return normalized


def _validate_positive_int(name: str, value: int, *, maximum: int) -> int:
    if value < 1:
        raise ValueError(f"{name} must be >= 1")
    if value > maximum:
        raise ValueError(f"{name} must be <= {maximum}")
    return value


def _domain_from_url(url: str) -> str:
    return urlparse(url).hostname or ""


def _validate_urls(urls: list[str], *, maximum: int) -> list[str]:
    if not urls:
        raise ValueError("urls must not be empty")
    if len(urls) > maximum:
        raise ValueError(f"urls must contain at most {maximum} entries")
    normalized: list[str] = []
    for url in urls:
        value = url.strip()
        parsed = urlparse(value)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            raise ValueError(f"invalid URL: {url!r}")
        normalized.append(value)
    return normalized


async def _maybe_await(value):
    if isawaitable(value):
        return await value
    return value


async def _ctx_get_state(ctx: Context, key: str):
    return await _maybe_await(ctx.get_state(key))


async def _ctx_set_state(ctx: Context, key: str, value) -> None:
    await _maybe_await(ctx.set_state(key, value))


def _normalize_domains(domains: list[str] | None, *, field_name: str) -> list[str]:
    if not domains:
        return []
    normalized: list[str] = []
    seen: set[str] = set()
    for domain in domains:
        value = domain.strip().lower()
        if not value:
            continue
        if value.startswith("www."):
            value = value[4:]
        if "/" in value:
            raise ValueError(f"{field_name} entries must be bare domains, got {domain!r}")
        if value not in seen:
            seen.add(value)
            normalized.append(value)
    return normalized


def _normalize_glob_patterns(patterns: list[str] | None, *, field_name: str) -> list[str]:
    if not patterns:
        return []
    normalized: list[str] = []
    for pattern in patterns:
        value = pattern.strip()
        if not value:
            continue
        normalized.append(value)
    if not normalized:
        raise ValueError(f"{field_name} must not be empty when provided")
    return normalized


def _normalize_mode(mode: str | None) -> str:
    normalized = (mode or "balanced").strip().lower()
    if normalized not in VALID_MODES:
        raise ValueError(f"invalid mode: {mode!r}. Expected one of {sorted(VALID_MODES)}")
    return normalized


def _normalize_categories(categories: list[str] | None) -> list[str] | None:
    if not categories:
        return None
    normalized: list[str] = []
    for cat in categories:
        value = cat.strip().lower()
        if value not in VALID_CATEGORIES:
            raise ValueError(f"invalid category: {cat!r}. Expected one of {sorted(VALID_CATEGORIES)}")
        normalized.append(value)
    return sorted(set(normalized)) or None


def _normalize_safesearch(safesearch: int | None) -> int | None:
    if safesearch is None:
        return None
    if safesearch not in {0, 1, 2}:
        raise ValueError(f"invalid safesearch: {safesearch!r}. Expected 0, 1, or 2")
    return safesearch


def _match_domain(domain: str, patterns: list[str]) -> bool:
    return any(domain == pattern or domain.endswith(f".{pattern}") for pattern in patterns)


def _filter_results_by_domain(
    results: list[dict],
    include_domains: list[str],
    exclude_domains: list[str],
) -> list[dict]:
    filtered: list[dict] = []
    for result in results:
        domain = _domain_from_url(result.get("url", "")).lower()
        bare_domain = domain[4:] if domain.startswith("www.") else domain
        if include_domains and not _match_domain(bare_domain, include_domains):
            continue
        if exclude_domains and _match_domain(bare_domain, exclude_domains):
            continue
        filtered.append(result)
    return filtered


def _diversify_ranked_entries(ranked_entry_idxs: list[int], entries: list[dict]) -> list[int]:
    """Interleave domains so the top results are not monopolized by one source."""
    by_domain: dict[str, list[int]] = defaultdict(list)
    domain_order: list[str] = []
    for eidx in ranked_entry_idxs:
        domain = _domain_from_url(entries[eidx]["url"]).lower() or entries[eidx]["url"]
        if domain not in by_domain:
            domain_order.append(domain)
        by_domain[domain].append(eidx)

    diversified: list[int] = []
    while by_domain:
        next_round: list[str] = []
        for domain in domain_order:
            queue = by_domain.get(domain)
            if not queue:
                continue
            diversified.append(queue.pop(0))
            if queue:
                next_round.append(domain)
            else:
                by_domain.pop(domain, None)
        domain_order = next_round
    return diversified


def _chunk_text(text: str) -> list[str]:
    """Split text into paragraph-based chunks suitable for reranking.

    Keeps paragraphs intact. Only splits paragraphs longer than 2000 chars
    by sentence boundaries so the reranker can score them effectively.
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    for para in paragraphs:
        if len(para) <= 2000:
            chunks.append(para)
            continue
        # Long paragraph — split by sentence into ~1000-char groups.
        sentences = _SENTENCE_SPLITTER.segment(para)
        current = ""
        for sent in sentences:
            if current and len(current) + len(sent) + 1 > 1000:
                chunks.append(current)
                current = sent
            else:
                current = f"{current} {sent}".strip() if current else sent
        if current:
            chunks.append(current)
    return chunks[:_MAX_CHUNKS_PER_PAGE]


async def _search(
    query: str,
    num_results: int = 10,
    time_range: str | None = None,
    categories: list[str] | None = None,
    language: str | None = None,
    safesearch: int | None = None,
    pageno: int = 1,
) -> list[dict]:
    """Query SearXNG and return raw results."""
    params: dict = {
        "q": query,
        "format": "json",
        "number_of_results": num_results,
    }
    if time_range:
        params["time_range"] = time_range.strip('"')
    if categories:
        params["categories"] = ",".join(categories)
    if language:
        params["language"] = language.strip()
    if safesearch is not None:
        params["safesearch"] = safesearch
    if pageno > 1:
        params["pageno"] = pageno

    async with httpx.AsyncClient(timeout=_HTTP_TIMEOUT) as client:
        resp = await client.get(f"{SEARXNG_URL}/search", params=params)
        resp.raise_for_status()
        return resp.json().get("results", [])[:num_results]


def _mode_settings(mode: str, num_results: int, scrape_top: int) -> tuple[int, int]:
    if mode == "deep":
        candidate_count = max(num_results * 2, 20)
        deep_scrape_top = max(scrape_top, min(num_results, MAX_SCRAPE))
        return candidate_count, min(deep_scrape_top, MAX_SCRAPE)
    return num_results, scrape_top


async def _probe_dependency(url: str) -> dict[str, str]:
    try:
        async with httpx.AsyncClient(timeout=3) as client:
            resp = await client.get(url)
            resp.raise_for_status()
        return {"status": "ok"}
    except httpx.HTTPError as exc:
        return {"status": "error", "detail": str(exc)}


def _extract_markdown(result: dict) -> str | None:
    html = result.get("html")
    if html:
        try:
            extracted = trafilatura.extract(
                html, output_format="txt", include_links=True, include_tables=True,
            )
            if extracted and len(extracted.strip()) >= 50:
                return extracted
        except Exception:
            pass

    md = result.get("markdown")
    if isinstance(md, dict):
        return md.get("fit_markdown") or md.get("raw_markdown")
    if isinstance(md, str):
        return md
    return result.get("cleaned_html")


def _content_type_without_charset(content_type: str | None) -> str | None:
    if not content_type:
        return None
    return content_type.split(";", 1)[0].strip().lower() or None


def _guess_file_type(url: str, content_type: str | None) -> str:
    normalized_content_type = _content_type_without_charset(content_type)
    suffix = (urlparse(url).path.rsplit(".", 1)[-1].lower() if "." in urlparse(url).path else "")

    if normalized_content_type == "application/pdf" or suffix == "pdf":
        return "pdf"
    if (
        normalized_content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == "docx"
    ):
        return "docx"
    if normalized_content_type in {"text/html", "application/xhtml+xml"} or suffix in {"html", "htm", "xhtml"}:
        return "html"
    if normalized_content_type in {"text/markdown", "text/x-markdown"} or suffix == "md":
        return "markdown"
    if normalized_content_type == "application/json" or suffix == "json":
        return "json"
    if normalized_content_type in {"application/xml", "text/xml"} or suffix in {"xml", "rss", "atom"}:
        return "xml"
    if normalized_content_type in {"text/csv", "application/csv", "application/vnd.ms-excel"} or suffix == "csv":
        return "csv"
    if normalized_content_type and normalized_content_type.startswith("text/"):
        return "text"
    guessed_type, _ = mimetypes.guess_type(url)
    guessed_type = _content_type_without_charset(guessed_type)
    if guessed_type and guessed_type.startswith("text/"):
        return "text"
    return "unknown"


def _extract_crawl_result(data: dict) -> dict:
    results = data.get("results")
    if results and isinstance(results, list):
        return results[0]
    return data.get("result", data)


def _extract_crawl_title(result: dict) -> str | None:
    metadata = result.get("metadata")
    if isinstance(metadata, dict):
        title = metadata.get("title")
        if title:
            return title
    title = result.get("title")
    if isinstance(title, str) and title.strip():
        return title.strip()
    return None


def _extract_crawl_links(result: dict, base_url: str) -> list[dict]:
    link_groups = result.get("links")
    if not isinstance(link_groups, dict):
        return []

    links: list[dict] = []
    for link_type in ("internal", "external"):
        bucket = link_groups.get(link_type) or []
        if not isinstance(bucket, list):
            continue
        for link in bucket:
            if not isinstance(link, dict):
                continue
            href = link.get("href")
            if not isinstance(href, str) or not href.strip():
                continue
            absolute_url = urljoin(base_url, href.strip())
            parsed = urlparse(absolute_url)
            if parsed.scheme not in {"http", "https"} or not parsed.netloc:
                continue
            links.append({
                "url": absolute_url,
                "title": (link.get("title") or "").strip() or None,
                "text": (link.get("text") or "").strip() or None,
                "link_type": link_type,
            })
    return links


def _build_crawl_config(
    js_code: list[str] | None = None,
    wait_for: str | None = None,
    page_timeout: int | None = None,
    screenshot: bool = False,
    remove_overlays: bool = True,
    scroll_full_page: bool = False,
) -> dict:
    """Build a CrawlerRunConfig payload for Crawl4AI's /crawl endpoint."""
    params: dict = {
        "excluded_tags": ["nav", "footer", "header", "aside"],
        "markdown_generator": {
            "type": "DefaultMarkdownGenerator",
            "params": {
                "content_filter": {
                    "type": "PruningContentFilter",
                    "params": {
                        "threshold": 0.48,
                        "threshold_type": "fixed",
                        "min_word_threshold": 0,
                    },
                },
            },
        },
    }
    if js_code:
        params["js_code"] = js_code
    if wait_for:
        params["wait_for"] = wait_for
    if page_timeout is not None:
        params["page_timeout"] = page_timeout
    if screenshot:
        params["screenshot"] = True
    if remove_overlays:
        params["remove_overlay_elements"] = True
    if scroll_full_page:
        params["scan_full_page"] = True
    return {"type": "CrawlerRunConfig", "params": params}


_DEFAULT_CRAWL_CONFIG = _build_crawl_config()


async def _scrape_impl(url: str, crawl_config: dict | None = None) -> dict:
    """Scrape a URL via Crawl4AI. Returns {content, title, screenshot}."""
    config = crawl_config or _DEFAULT_CRAWL_CONFIG
    async with httpx.AsyncClient(timeout=_HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{CRAWL4AI_URL}/crawl",
            json={"urls": [url], "priority": 8, "crawler_config": config},
        )
        resp.raise_for_status()
        data = resp.json()

        task_id = data.get("task_id")
        if task_id:
            while True:
                await asyncio.sleep(1)
                status_resp = await client.get(f"{CRAWL4AI_URL}/task/{task_id}")
                status_resp.raise_for_status()
                status_data = status_resp.json()
                status = status_data.get("status")
                if status == "completed":
                    result = status_data.get("result", {})
                    return {
                        "content": _extract_markdown(result),
                        "title": _extract_crawl_title(result),
                        "screenshot": result.get("screenshot"),
                    }
                if status == "failed":
                    return {"content": None, "title": None, "screenshot": None}

        result = _extract_crawl_result(data)
        return {
            "content": _extract_markdown(result),
            "title": _extract_crawl_title(result),
            "screenshot": result.get("screenshot"),
        }


async def _scrape(url: str, crawl_config: dict | None = None) -> dict:
    """Scrape a URL via Crawl4AI, bounded by REQUEST_TIMEOUT seconds end-to-end.

    Returns {content, title, screenshot}. On failure content is None.
    """
    empty = {"content": None, "title": None, "screenshot": None}
    try:
        return await asyncio.wait_for(_scrape_impl(url, crawl_config), timeout=REQUEST_TIMEOUT)
    except asyncio.TimeoutError:
        log.warning("scrape timed out url=%s budget=%ss", url, REQUEST_TIMEOUT)
    except httpx.HTTPError as e:
        log.warning("scrape http error url=%s err=%s", url, e)
    except (ValueError, KeyError) as e:
        log.warning("scrape payload error url=%s err=%s", url, e)
    return empty


async def _scrape_cached(
    url: str, cache: dict[str, str | None], crawl_config: dict | None = None,
) -> str | None:
    """Scrape with per-session cache. Returns cached content on hit, scrapes on miss.

    Bypasses cache when a custom crawl_config is provided (browser actions are intentionally fresh).
    """
    if crawl_config is None and url in cache:
        log.debug("scrape cache hit url=%s", url)
        return cache[url]
    result = await _scrape(url, crawl_config)
    content = result["content"]
    cache[url] = content
    return content


async def _head_content_type(url: str) -> str | None:
    try:
        async with httpx.AsyncClient(timeout=_HTTP_TIMEOUT, follow_redirects=True) as client:
            resp = await client.head(url)
            resp.raise_for_status()
            return resp.headers.get("content-type")
    except httpx.HTTPError:
        return None


async def _looks_like_pdf(url: str) -> bool:
    if urlparse(url).path.lower().endswith(".pdf"):
        return True
    content_type = await _head_content_type(url)
    return bool(content_type and "application/pdf" in content_type.lower())


async def _detect_file_type(url: str) -> tuple[str, str | None]:
    content_type = await _head_content_type(url)
    return _guess_file_type(url, content_type), _content_type_without_charset(content_type)


def _pdf_bytes_to_markdown(
    data: bytes,
    start_page: int | None = None,
    end_page: int | None = None,
) -> tuple[str, str | None, int, int]:
    """Extract PDF pages to markdown.

    Returns (content, title, total_pages, last_page_included).
    Pages are added incrementally; if the budget is exhausted mid-document
    the returned last_page_included reflects where extraction actually stopped.
    """
    reader = PdfReader(io.BytesIO(data))
    total_pages = len(reader.pages)
    title = None
    if reader.metadata:
        title = reader.metadata.title
    first = (start_page or 1) - 1
    last = min(end_page or total_pages, total_pages)
    sections: list[str] = []
    char_budget = _MAX_CONTENT_CHARS
    last_page_included = first  # 0-indexed; will convert to 1-indexed at return
    for page_number, page in enumerate(reader.pages[first:last], start=first + 1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        section = f"## Page {page_number}\n\n{text}"
        if char_budget - len(section) < 0 and sections:
            break
        sections.append(section)
        char_budget -= len(section) + 2  # account for "\n\n" join separator
        last_page_included = page_number
    return "\n\n".join(sections), title, total_pages, last_page_included


def _docx_bytes_to_markdown(data: bytes) -> tuple[str, str | None]:
    document = DocxDocument(io.BytesIO(data))
    title = document.core_properties.title or None
    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(text)

    for table_number, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            sections.append(f"## Table {table_number}\n\n" + "\n".join(rows))

    return "\n\n".join(sections)[:_MAX_CONTENT_CHARS], title


async def _extract_pdf_document(
    url: str,
    start_page: int | None = None,
    end_page: int | None = None,
) -> dict:
    async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT, follow_redirects=True) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        content, title, total_pages, last_page = _pdf_bytes_to_markdown(
            resp.content, start_page=start_page, end_page=end_page,
        )
        return {
            "status": "ok",
            "url": url,
            "content_type": "application/pdf",
            "file_type": "pdf",
            "title": title,
            "content": content,
            "total_pages": total_pages,
            "start_page": start_page or 1,
            "end_page": last_page,
        }


async def _extract_web_document(url: str, crawl_config: dict | None = None) -> dict:
    result = await _scrape(url, crawl_config)
    content = result["content"]
    if not content:
        return {
            "status": "error",
            "url": url,
            "content_type": "text/html",
            "file_type": "html",
            "title": None,
            "content": "",
            "error": "extraction failed",
        }
    doc = {
        "status": "ok",
        "url": url,
        "content_type": "text/html",
        "file_type": "html",
        "title": result.get("title"),
        "content": content[:_MAX_CONTENT_CHARS],
    }
    if result.get("screenshot"):
        doc["screenshot"] = result["screenshot"]
    return doc


async def _extract_docx_document(url: str) -> dict:
    async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT, follow_redirects=True) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        content, title = _docx_bytes_to_markdown(resp.content)
        return {
            "status": "ok",
            "url": url,
            "content_type": _content_type_without_charset(resp.headers.get("content-type"))
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_type": "docx",
            "title": title,
            "content": content,
        }


async def _extract_text_document(url: str, file_type: str) -> dict:
    async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT, follow_redirects=True) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        return {
            "status": "ok",
            "url": url,
            "content_type": _content_type_without_charset(resp.headers.get("content-type"))
            or mimetypes.guess_type(url)[0]
            or "text/plain",
            "file_type": file_type,
            "title": None,
            "content": resp.text[:_MAX_CONTENT_CHARS],
        }


async def _discover_page_links(url: str) -> dict:
    async with httpx.AsyncClient(timeout=_HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{CRAWL4AI_URL}/crawl",
            json={"urls": [url], "priority": 6, "crawler_config": _DEFAULT_CRAWL_CONFIG},
        )
        resp.raise_for_status()
        data = resp.json()

        task_id = data.get("task_id")
        if task_id:
            while True:
                await asyncio.sleep(1)
                status_resp = await client.get(f"{CRAWL4AI_URL}/task/{task_id}")
                status_resp.raise_for_status()
                status_data = status_resp.json()
                status = status_data.get("status")
                if status == "completed":
                    result = _extract_crawl_result(status_data.get("result", {}))
                    return {
                        "status": "ok",
                        "url": url,
                        "title": _extract_crawl_title(result),
                        "links": _extract_crawl_links(result, url),
                    }
                if status == "failed":
                    return {
                        "status": "error",
                        "url": url,
                        "title": None,
                        "links": [],
                        "error": "link discovery failed",
                    }

        result = _extract_crawl_result(data)
        return {
            "status": "ok",
            "url": url,
            "title": _extract_crawl_title(result),
            "links": _extract_crawl_links(result, url),
        }


def _url_matches_patterns(url: str, patterns: list[str]) -> bool:
    return any(fnmatch.fnmatch(url, pattern) for pattern in patterns)


def _rank_document_content(query: str | None, content: str) -> tuple[str, list[dict]]:
    if not query or not content:
        return content[:_MAX_CONTENT_CHARS], []
    chunks = _chunk_text(content[:_MAX_CONTENT_CHARS])
    if not chunks:
        return content[:_MAX_CONTENT_CHARS], []
    scored = _rerank_scored(query, chunks)
    top = [{"text": chunks[idx], "score": score} for idx, score in scored[:_TOP_CHUNKS]]
    if not top:
        return content[:_MAX_CONTENT_CHARS], []
    return "\n\n".join(item["text"] for item in top), top


async def _extract_url_document(
    url: str,
    query: str | None,
    cache: dict[str, dict],
    crawl_config: dict | None = None,
    start_page: int | None = None,
    end_page: int | None = None,
) -> dict:
    # Skip cache when a page range is specified (caller wants a specific window).
    if crawl_config is None and start_page is None and end_page is None and url in cache:
        cached = cache[url]
        content, top_chunks = _rank_document_content(query, cached.get("content", ""))
        return {
            **cached,
            "content": content,
            "top_chunks": top_chunks,
            "cached": True,
        }

    file_type = "unknown"
    content_type = None
    try:
        file_type, content_type = await _detect_file_type(url)
        if file_type == "pdf":
            extracted = await _extract_pdf_document(url, start_page=start_page, end_page=end_page)
        elif file_type == "docx":
            extracted = await _extract_docx_document(url)
        elif file_type in {"text", "markdown", "json", "xml", "csv"}:
            extracted = await _extract_text_document(url, file_type)
        else:
            extracted = await _extract_web_document(url, crawl_config)
    except Exception as exc:
        extracted = {
            "status": "error",
            "url": url,
            "content_type": content_type,
            "file_type": file_type,
            "title": None,
            "content": "",
            "error": str(exc),
        }

    if extracted["status"] == "ok":
        cache[url] = {
            "status": extracted["status"],
            "url": url,
            "content_type": extracted.get("content_type"),
            "title": extracted.get("title"),
            "content": extracted.get("content", ""),
        }
        content, top_chunks = _rank_document_content(query, extracted.get("content", ""))
        extracted["content"] = content
        extracted["top_chunks"] = top_chunks
        extracted["cached"] = False
        return extracted

    extracted["top_chunks"] = []
    extracted["cached"] = False
    return extracted


log.info("loading reranker model=%s max_length=%d", RERANK_MODEL, _RERANK_MAX_LENGTH)
_ranker = Ranker(model_name=RERANK_MODEL, max_length=_RERANK_MAX_LENGTH)
log.info("reranker ready")


def _rerank_scored(query: str, documents: list[str]) -> list[tuple[int, float]]:
    """Rerank documents via FlashRank. Returns (index, score) pairs sorted by descending relevance."""
    if not documents:
        return []
    passages = [{"id": i, "text": doc, "meta": {}} for i, doc in enumerate(documents)]
    request = RerankRequest(query=query, passages=passages)
    results = _ranker.rerank(request)
    return [(r["id"], float(r["score"])) for r in results]


@mcp.custom_route("/health", methods=["GET"])
async def health(_: Request) -> JSONResponse:
    return JSONResponse({"status": "ok", "reranker": {"name": "flashrank", "model": RERANK_MODEL}})


@mcp.custom_route("/ready", methods=["GET"])
async def ready(_: Request) -> JSONResponse:
    searxng = await _probe_dependency(f"{SEARXNG_URL}/healthz")
    crawl4ai = await _probe_dependency(f"{CRAWL4AI_URL}/health")
    ready_ok = searxng["status"] == "ok" and crawl4ai["status"] == "ok"
    payload = {
        "status": "ok" if ready_ok else "degraded",
        "dependencies": {
            "searxng": searxng,
            "crawl4ai": crawl4ai,
            "reranker": {"status": "ok", "name": "flashrank", "model": RERANK_MODEL},
        },
    }
    return JSONResponse(payload, status_code=200 if ready_ok else 503)



async def _web_search_impl(
    query: str,
    num_results: int = 10,
    scrape_top: int = MAX_SCRAPE,
    time_range: str | None = None,
    mode: str = "balanced",
    include_domains: list[str] | None = None,
    exclude_domains: list[str] | None = None,
    categories: list[str] | None = None,
    language: str | None = None,
    safesearch: int | None = None,
    ctx: Context | None = None,
) -> dict:
    """Search the web, scrape top results, and return structured JSON ranked by relevance.

    Pipeline: SearXNG search -> Crawl4AI scrape -> chunk -> FlashRank reranker -> formatted output.
    Scraped pages are split into paragraphs and reranked at the chunk level, so only
    the most query-relevant excerpts from each page are returned.
    Results are cached within the session.

    Args:
        query: The search query. Do not include dates or years in the query — use time_range instead.
        num_results: Number of search results to fetch (default 10).
        scrape_top: Number of top results to scrape for full content (default 5).
        time_range: Time filter for recency: 'day', 'week', 'month', or 'year'. Use this instead of adding dates to the query. Omit or pass null for no filter.
    """
    query = _validate_query(query)
    num_results = _validate_positive_int("num_results", num_results, maximum=MAX_RESULTS)
    scrape_top = _validate_positive_int("scrape_top", scrape_top, maximum=MAX_SCRAPE)
    time_range = _normalize_time_range(time_range)
    mode = _normalize_mode(mode)
    include_domains = _normalize_domains(include_domains, field_name="include_domains")
    exclude_domains = _normalize_domains(exclude_domains, field_name="exclude_domains")
    categories = _normalize_categories(categories)
    language = language.strip() if language else None
    safesearch = _normalize_safesearch(safesearch)
    candidate_count, scrape_top = _mode_settings(mode, num_results, scrape_top)
    started = time.monotonic()
    warnings: list[dict] = []
    degraded = False
    timings_ms = {"search": 0, "scrape": 0, "rerank": 0, "total": 0}

    # --- session state ---
    scrape_cache: dict[str, str | None] = {}
    query_cache: dict[str, dict] = {}
    seen_urls: set[str] = set()
    if ctx:
        scrape_cache = await _ctx_get_state(ctx, STATE_SCRAPE_CACHE) or {}
        query_cache = await _ctx_get_state(ctx, STATE_QUERY_CACHE) or {}
        seen_urls = set(await _ctx_get_state(ctx, STATE_SEEN_URLS) or [])

    # --- exact query cache ---
    qkey = hashlib.sha256(
        json.dumps(
            [
                query.lower().strip(),
                num_results,
                scrape_top,
                time_range,
                mode,
                include_domains,
                exclude_domains,
                categories,
                language,
                safesearch,
            ],
            sort_keys=True,
        ).encode()
    ).hexdigest()
    if qkey in query_cache:
        log.info("query cache hit query=%r", query)
        return query_cache[qkey]

    # --- search ---
    search_started = time.monotonic()
    try:
        results = await _search(
            query,
            num_results=candidate_count,
            time_range=time_range,
            categories=categories,
            language=language,
            safesearch=safesearch,
        )
    except Exception as exc:
        degraded = True
        warnings.append(_warning("search_failed", "searxng", str(exc)))
        results = []

    if results and mode == "deep":
        try:
            page2 = await _search(
                query,
                num_results=candidate_count,
                time_range=time_range,
                categories=categories,
                language=language,
                safesearch=safesearch,
                pageno=2,
            )
            results.extend(page2)
        except Exception as exc:
            warnings.append(_warning("search_failed", "searxng", f"page 2: {exc}"))

    timings_ms["search"] = int((time.monotonic() - search_started) * 1000)
    if results:
        results = _filter_results_by_domain(results, include_domains, exclude_domains)
    if not results:
        response = {
            "query": query,
            "time_range": time_range,
            "mode": mode,
            "include_domains": include_domains,
            "exclude_domains": exclude_domains,
            "categories": categories,
            "language": language,
            "safesearch": safesearch,
            "results": [],
            "meta": {
                "num_results_requested": num_results,
                "num_results_returned": 0,
                "scrape_top": min(scrape_top, num_results),
                "candidate_count": candidate_count,
                "search_backend": "searxng",
                "reranker": {"name": "flashrank", "model": RERANK_MODEL},
                "degraded": degraded,
                "warnings": warnings or [_warning("no_results", "searxng", query)],
                "timings_ms": {
                    **timings_ms,
                    "total": int((time.monotonic() - started) * 1000),
                },
            },
        }
        query_cache[qkey] = response
        if ctx:
            await _ctx_set_state(ctx, STATE_QUERY_CACHE, query_cache)
        return response
    results = _dedup_results(results)
    results = results[:candidate_count]

    # --- scrape (cache-aware) ---
    to_scrape = min(scrape_top, len(results))
    scrape_started = time.monotonic()
    scrape_tasks = [_scrape_cached(r["url"], scrape_cache) for r in results[:to_scrape]]
    scraped = await asyncio.gather(*scrape_tasks)
    timings_ms["scrape"] = int((time.monotonic() - scrape_started) * 1000)
    scrape_failures = sum(1 for content in scraped if content is None)
    if scrape_failures:
        degraded = True
        warnings.append(_warning("scrape_failed", "crawl4ai", f"{scrape_failures} of {to_scrape} pages failed"))

    # --- build entries ---
    entries: list[dict] = []
    for i, result in enumerate(results[:to_scrape]):
        raw = scraped[i][:_MAX_CONTENT_CHARS] if scraped[i] else None
        entries.append({
            "title": result.get("title", "Untitled"),
            "url": result.get("url", ""),
            "content": raw or result.get("content", ""),
            "scraped": raw is not None,
        })

    for result in results[to_scrape:]:
        entries.append({
            "title": result.get("title", "Untitled"),
            "url": result.get("url", ""),
            "content": result.get("content", ""),
            "scraped": False,
        })

    # --- chunk scraped pages, keep snippets as single chunks ---
    all_chunks: list[str] = []
    chunk_to_entry: list[int] = []
    for i, entry in enumerate(entries):
        if entry["scraped"] and entry["content"]:
            chunks = _chunk_text(entry["content"])
        else:
            chunks = [entry["content"]] if entry["content"] else []
        for chunk in chunks:
            all_chunks.append(chunk)
            chunk_to_entry.append(i)

    # --- deduplicate near-identical chunks across pages ---
    all_chunks, chunk_to_entry = _dedup_chunks(all_chunks, chunk_to_entry)

    # --- rerank at the chunk level ---
    rerank_started = time.monotonic()
    rerank_failed = False
    try:
        scored = _rerank_scored(query, all_chunks)
    except Exception as exc:
        log.warning("rerank failed query=%r err=%s", query, exc)
        warnings.append(_warning("rerank_failed", "flashrank", str(exc)))
        degraded = True
        rerank_failed = True
        scored = []
    timings_ms["rerank"] = int((time.monotonic() - rerank_started) * 1000)

    # --- group scores by entry, select top-K chunks per page ---
    entry_chunks: dict[int, list[tuple[str, float]]] = defaultdict(list)
    for chunk_idx, score in scored:
        eidx = chunk_to_entry[chunk_idx]
        entry_chunks[eidx].append((all_chunks[chunk_idx], score))

    for eidx in entry_chunks:
        entry_chunks[eidx].sort(key=lambda x: x[1], reverse=True)
        entry_chunks[eidx] = entry_chunks[eidx][:_TOP_CHUNKS]

    # --- rank pages by best chunk score ---
    entry_best: dict[int, float | None] = {
        eidx: chunks[0][1] for eidx, chunks in entry_chunks.items() if chunks
    }
    if rerank_failed:
        ranked_entry_idxs = list(range(len(entries)))
    else:
        ranked_entry_idxs = sorted(entry_best, key=entry_best.get, reverse=True)
        for i in range(len(entries)):
            if i not in entry_best:
                ranked_entry_idxs.append(i)
    ranked_entry_idxs = _diversify_ranked_entries(ranked_entry_idxs, entries)

    # --- format structured output ---
    structured_results: list[dict] = []
    new_urls: list[str] = []
    for rank, eidx in enumerate(ranked_entry_idxs, 1):
        entry = entries[eidx]
        url = entry["url"]
        normalized_url = _normalize_url(url)
        top = entry_chunks.get(eidx, [])
        if top:
            content = "\n\n".join(chunk for chunk, _ in top)
        else:
            content = entry["content"]

        structured_results.append({
            "rank": rank,
            "search_rank": eidx + 1,
            "title": entry["title"],
            "url": url,
            "normalized_url": normalized_url,
            "domain": _domain_from_url(url),
            "snippet": results[eidx].get("content", ""),
            "content": content,
            "top_chunks": [{"text": chunk, "score": score} for chunk, score in top],
            "score": entry_best.get(eidx),
            "scraped": entry["scraped"],
            "previously_seen": normalized_url in seen_urls,
        })
        new_urls.append(normalized_url)

    response = {
        "query": query,
        "time_range": time_range,
        "mode": mode,
        "include_domains": include_domains,
        "exclude_domains": exclude_domains,
        "categories": categories,
        "language": language,
        "safesearch": safesearch,
        "results": structured_results,
        "meta": {
            "num_results_requested": num_results,
            "num_results_returned": len(structured_results),
            "scrape_top": to_scrape,
            "candidate_count": candidate_count,
            "search_backend": "searxng",
            "reranker": {"name": "flashrank", "model": RERANK_MODEL},
            "degraded": degraded,
            "warnings": warnings,
            "timings_ms": {
                **timings_ms,
                "total": int((time.monotonic() - started) * 1000),
            },
        },
    }

    # --- persist session state ---
    seen_urls.update(new_urls)
    query_cache[qkey] = response
    if ctx:
        await _ctx_set_state(ctx, STATE_SCRAPE_CACHE, scrape_cache)
        await _ctx_set_state(ctx, STATE_QUERY_CACHE, query_cache)
        await _ctx_set_state(ctx, STATE_SEEN_URLS, list(seen_urls))

    log.info(
        "query=%r chunks=%d pages=%d scrape_cache=%d query_cache=%d seen=%d",
        query, len(all_chunks), len(entries), len(scrape_cache), len(query_cache), len(seen_urls),
    )

    return response


# ---------------------------------------------------------------------------
# Markdown formatters – convert internal response dicts into LLM-friendly text
# ---------------------------------------------------------------------------

def _format_search_results(response: dict) -> str:
    """Format web_search/site_search response as markdown for LLM consumption."""
    parts = [f"query: {response['query']}"]
    if response.get("mode"):
        parts.append(f"mode: {response['mode']}")
    if response.get("time_range"):
        parts.append(f"time_range: {response['time_range']}")
    meta = response.get("meta", {})
    parts.append(f"results: {meta.get('num_results_returned', len(response.get('results', [])))}")
    if meta.get("degraded"):
        parts.append("status: degraded")
    warnings = meta.get("warnings", [])
    if warnings:
        parts.append("warnings: " + "; ".join(w.get("detail", str(w)) for w in warnings))
    header = "\n".join(parts)

    sections = [header, "---"]
    for r in response.get("results", []):
        title = r.get("title", "Untitled")
        url = r.get("url", "")
        content = r.get("content", "")
        section = f"## [{title}]({url})\n\n{content}" if content else f"## [{title}]({url})"
        sections.append(section)

    return "\n\n".join(sections)


def _format_image_results(response: dict) -> str:
    """Format image_search response as markdown."""
    parts = [f"query: {response['query']}"]
    meta = response.get("meta", {})
    parts.append(f"results: {meta.get('num_results_returned', len(response.get('results', [])))}")
    if meta.get("error"):
        parts.append(f"error: {meta['error']}")
    header = "\n".join(parts)

    sections = [header, "---"]
    for r in response.get("results", []):
        title = r.get("title", "Untitled")
        source_url = r.get("source_url", "")
        image_url = r.get("image_url", "")
        dims = r.get("dimensions")
        line = f"## [{title}]({source_url})\n\n![{title}]({image_url})"
        if dims:
            line += f"\n{dims}"
        sections.append(line)

    return "\n\n".join(sections)


def _format_extract_results(response: dict) -> str:
    """Format extract_url/extract_urls response as markdown."""
    parts = []
    if response.get("query"):
        parts.append(f"query: {response['query']}")
    # handle both single-result (extract_url) and multi-result (extract_urls)
    results = response.get("results", [])
    if "result" in response:
        results = [response["result"]]
    meta = response.get("meta", {})
    succeeded = meta.get("urls_succeeded", sum(1 for r in results if r.get("status") == "ok"))
    failed = meta.get("urls_failed", sum(1 for r in results if r.get("status") != "ok"))
    parts.append(f"succeeded: {succeeded}")
    if failed:
        parts.append(f"failed: {failed}")
    header = "\n".join(parts)

    sections = [header, "---"] if parts else ["---"]
    for r in results:
        title = r.get("title") or "Untitled"
        url = r.get("url", "")
        content = r.get("content", "")
        status = r.get("status", "")
        if status == "error":
            error = r.get("error", "extraction failed")
            section = f"## [{title}]({url})\n\n**Error:** {error}"
        elif content:
            section = f"## [{title}]({url})\n\n{content}"
        else:
            section = f"## [{title}]({url})"
        if r.get("total_pages") is not None:
            section += f"\n\n_pages {r['start_page']}-{r['end_page']} of {r['total_pages']}_"
        sections.append(section)

    return "\n\n".join(sections)


def _format_map_results(response: dict) -> str:
    """Format map_site response as markdown."""
    parts = [f"url: {response['url']}"]
    meta = response.get("meta", {})
    parts.append(f"urls_found: {meta.get('urls_returned', len(response.get('results', [])))}")
    warnings = meta.get("warnings", [])
    if warnings:
        parts.append("warnings: " + "; ".join(w.get("detail", str(w)) for w in warnings))
    header = "\n".join(parts)

    sections = [header, "---"]
    for r in response.get("results", []):
        title = r.get("title") or r.get("link_text") or r.get("url", "")
        url = r.get("url", "")
        depth = r.get("depth", 0)
        indent = "  " * depth
        sections.append(f"{indent}- [{title}]({url})")

    return "\n\n".join(sections)


def _format_crawl_results(response: dict) -> str:
    """Format crawl_site response as markdown."""
    parts = [f"url: {response['url']}"]
    if response.get("query"):
        parts.append(f"query: {response['query']}")
    meta = response.get("meta", {})
    parts.append(f"succeeded: {meta.get('urls_succeeded', 0)}")
    if meta.get("urls_failed"):
        parts.append(f"failed: {meta['urls_failed']}")
    warnings = meta.get("warnings", [])
    if warnings:
        parts.append("warnings: " + "; ".join(w.get("detail", str(w)) for w in warnings))
    header = "\n".join(parts)

    sections = [header, "---"]
    for r in response.get("results", []):
        title = r.get("title") or "Untitled"
        url = r.get("url", "")
        content = r.get("content", "")
        status = r.get("status", "")
        if status == "error":
            error = r.get("error", "extraction failed")
            section = f"## [{title}]({url})\n\n**Error:** {error}"
        elif content:
            section = f"## [{title}]({url})\n\n{content}"
        else:
            section = f"## [{title}]({url})"
        sections.append(section)

    return "\n\n".join(sections)


def _format_similar_results(response: dict) -> str:
    """Format find_similar response as markdown."""
    parts = [f"source: [{response.get('source_title') or 'Untitled'}]({response['source_url']})"]
    meta = response.get("meta", {})
    parts.append(f"results: {len(response.get('results', []))}")
    if meta.get("error"):
        parts.append(f"error: {meta['error']}")
    warnings = meta.get("warnings", [])
    if warnings:
        parts.append("warnings: " + "; ".join(w.get("detail", str(w)) for w in warnings))
    header = "\n".join(parts)

    sections = [header, "---"]
    for r in response.get("results", []):
        title = r.get("title", "Untitled")
        url = r.get("url", "")
        snippet = r.get("snippet", "")
        section = f"## [{title}]({url})\n\n{snippet}" if snippet else f"## [{title}]({url})"
        sections.append(section)

    return "\n\n".join(sections)


@mcp.tool
async def web_search(
    query: str,
    num_results: int = 10,
    scrape_top: int = MAX_SCRAPE,
    time_range: str | None = None,
    mode: str = "balanced",
    include_domains: list[str] | None = None,
    exclude_domains: list[str] | None = None,
    categories: list[str] | None = None,
    language: str | None = None,
    safesearch: int | None = None,
    ctx: Context | None = None,
) -> str:
    response = await _web_search_impl(
        query=query,
        num_results=num_results,
        scrape_top=scrape_top,
        time_range=time_range,
        mode=mode,
        include_domains=include_domains,
        exclude_domains=exclude_domains,
        categories=categories,
        language=language,
        safesearch=safesearch,
        ctx=ctx,
    )
    return _format_search_results(response)


@mcp.tool
async def image_search(
    query: str,
    num_results: int = 10,
    language: str | None = None,
    safesearch: int | None = None,
    time_range: str | None = None,
    ctx: Context | None = None,
) -> str:
    """Search for images via SearXNG. Returns image URLs, thumbnails, and source metadata."""
    query = _validate_query(query)
    num_results = _validate_positive_int("num_results", num_results, maximum=MAX_RESULTS)
    time_range = _normalize_time_range(time_range)
    language = language.strip() if language else None
    safesearch = _normalize_safesearch(safesearch)
    started = time.monotonic()

    try:
        raw = await _search(
            query,
            num_results=num_results,
            time_range=time_range,
            categories=["images"],
            language=language,
            safesearch=safesearch,
        )
    except Exception as exc:
        response = {
            "query": query,
            "results": [],
            "meta": {
                "num_results_requested": num_results,
                "num_results_returned": 0,
                "search_backend": "searxng",
                "error": str(exc),
                "timings_ms": {"total": int((time.monotonic() - started) * 1000)},
            },
        }
        return _format_image_results(response)

    results: list[dict] = []
    for rank, item in enumerate(raw, 1):
        results.append({
            "rank": rank,
            "title": item.get("title", ""),
            "image_url": item.get("img_src", ""),
            "thumbnail_url": item.get("thumbnail_src", ""),
            "source_url": item.get("url", ""),
            "source_domain": _domain_from_url(item.get("url", "")),
            "dimensions": item.get("resolution"),
            "format": item.get("img_format"),
        })

    response = {
        "query": query,
        "results": results,
        "meta": {
            "num_results_requested": num_results,
            "num_results_returned": len(results),
            "search_backend": "searxng",
            "timings_ms": {"total": int((time.monotonic() - started) * 1000)},
        },
    }
    return _format_image_results(response)


async def _extract_urls_impl(
    urls: list[str],
    query: str | None = None,
    crawl_config: dict | None = None,
    start_page: int | None = None,
    end_page: int | None = None,
    ctx: Context | None = None,
) -> dict:
    """Extract a batch of URLs with per-URL status reporting.

    Uses Crawl4AI for web pages and pypdf for PDF documents.
    Returns markdown content only in v1.
    """
    urls = _validate_urls(urls, maximum=_MAX_EXTRACT_URLS)
    normalized_query = query.strip() if query else None
    started = time.monotonic()

    extract_cache: dict[str, dict] = {}
    if ctx:
        extract_cache = await _ctx_get_state(ctx, STATE_EXTRACT_CACHE) or {}

    documents = await asyncio.gather(*[
        _extract_url_document(
            url, normalized_query, extract_cache, crawl_config,
            start_page=start_page, end_page=end_page,
        )
        for url in urls
    ])

    results: list[dict] = []
    urls_succeeded = 0
    urls_failed = 0
    for document in documents:
        if document["status"] == "ok":
            urls_succeeded += 1
        else:
            urls_failed += 1
        url = document["url"]
        entry = {
            "url": url,
            "normalized_url": _normalize_url(url),
            "domain": _domain_from_url(url),
            "status": document["status"],
            "content_type": document.get("content_type"),
            "file_type": document.get("file_type"),
            "title": document.get("title"),
            "content": document.get("content", ""),
            "top_chunks": document.get("top_chunks", []),
            "cached": document.get("cached", False),
            "error": document.get("error"),
        }
        if document.get("screenshot"):
            entry["screenshot"] = document["screenshot"]
        if document.get("total_pages") is not None:
            entry["total_pages"] = document["total_pages"]
            entry["start_page"] = document.get("start_page", 1)
            entry["end_page"] = document.get("end_page")
        results.append(entry)

    if ctx:
        await _ctx_set_state(ctx, STATE_EXTRACT_CACHE, extract_cache)

    response = {
        "query": normalized_query,
        "results": results,
        "meta": {
            "urls_requested": len(urls),
            "urls_succeeded": urls_succeeded,
            "urls_failed": urls_failed,
            "timings_ms": {
                "total": int((time.monotonic() - started) * 1000),
            },
        },
    }
    log.info("extract_urls requested=%d succeeded=%d failed=%d", len(urls), urls_succeeded, urls_failed)
    return response


def _maybe_build_crawl_config(
    js_code: list[str] | None,
    wait_for: str | None,
    page_timeout: int | None,
    screenshot: bool,
    remove_overlays: bool,
    scroll_full_page: bool,
) -> dict | None:
    """Build a custom crawl config only if any browser param is non-default."""
    if any([js_code, wait_for, page_timeout is not None, screenshot, not remove_overlays, scroll_full_page]):
        return _build_crawl_config(
            js_code=js_code,
            wait_for=wait_for,
            page_timeout=page_timeout,
            screenshot=screenshot,
            remove_overlays=remove_overlays,
            scroll_full_page=scroll_full_page,
        )
    return None


@mcp.tool
async def extract_url(
    url: str,
    query: str | None = None,
    start_page: int | None = None,
    end_page: int | None = None,
    js_code: list[str] | None = None,
    wait_for: str | None = None,
    page_timeout: int | None = None,
    screenshot: bool = False,
    remove_overlays: bool = True,
    scroll_full_page: bool = False,
    ctx: Context | None = None,
) -> str:
    """Extract content from a single URL. For PDFs, use start_page/end_page to read specific page ranges (1-indexed, inclusive). The response includes total_pages so you can paginate through large documents."""
    crawl_config = _maybe_build_crawl_config(
        js_code, wait_for, page_timeout, screenshot, remove_overlays, scroll_full_page,
    )
    response = await _extract_urls_impl(
        urls=[url], query=query, crawl_config=crawl_config,
        start_page=start_page, end_page=end_page, ctx=ctx,
    )
    result = response["results"][0]
    response_dict = {
        "query": response["query"],
        "result": result,
        "meta": {
            **response["meta"],
            "url": result["url"],
        },
    }
    return _format_extract_results(response_dict)


@mcp.tool
async def extract_urls(
    urls: list[str],
    query: str | None = None,
    start_page: int | None = None,
    end_page: int | None = None,
    js_code: list[str] | None = None,
    wait_for: str | None = None,
    page_timeout: int | None = None,
    screenshot: bool = False,
    remove_overlays: bool = True,
    scroll_full_page: bool = False,
    ctx: Context | None = None,
) -> str:
    """Extract content from multiple URLs. For PDFs, use start_page/end_page to read specific page ranges (1-indexed, inclusive). The page range applies to all PDF URLs in the batch."""
    crawl_config = _maybe_build_crawl_config(
        js_code, wait_for, page_timeout, screenshot, remove_overlays, scroll_full_page,
    )
    response = await _extract_urls_impl(
        urls=urls, query=query, crawl_config=crawl_config,
        start_page=start_page, end_page=end_page, ctx=ctx,
    )
    return _format_extract_results(response)


async def _map_site_impl(
    url: str,
    max_urls: int = 25,
    max_depth: int = 1,
    include_patterns: list[str] | None = None,
    exclude_patterns: list[str] | None = None,
    same_domain_only: bool = True,
) -> dict:
    """Discover in-scope URLs from a site using Crawl4AI link extraction."""
    root_url = _validate_urls([url], maximum=1)[0]
    max_urls = _validate_positive_int("max_urls", max_urls, maximum=_MAX_MAP_URLS)
    max_depth = _validate_positive_int("max_depth", max_depth, maximum=_MAX_MAP_DEPTH)
    include_patterns = _normalize_glob_patterns(include_patterns, field_name="include_patterns")
    exclude_patterns = _normalize_glob_patterns(exclude_patterns, field_name="exclude_patterns")

    started = time.monotonic()
    root_domain = _domain_from_url(root_url).lower()
    bare_root_domain = root_domain[4:] if root_domain.startswith("www.") else root_domain

    queue = deque([(root_url, 0, None)])
    visited_pages: set[str] = set()
    discovered: dict[str, dict] = {}
    warnings: list[dict] = []

    while queue and len(discovered) < max_urls:
        current_url, depth, discovered_from = queue.popleft()
        normalized_current = _normalize_url(current_url)
        if normalized_current in visited_pages:
            continue
        visited_pages.add(normalized_current)

        current_entry = discovered.setdefault(normalized_current, {
            "url": current_url,
            "normalized_url": normalized_current,
            "domain": _domain_from_url(current_url),
            "title": None,
            "link_text": None,
            "depth": depth,
            "discovered_from": discovered_from,
            "link_type": "seed" if depth == 0 else "internal",
        })

        if depth >= max_depth:
            continue

        try:
            page = await asyncio.wait_for(_discover_page_links(current_url), timeout=REQUEST_TIMEOUT)
        except asyncio.TimeoutError:
            warnings.append(_warning("link_discovery_timeout", "crawl4ai", current_url))
            continue
        except httpx.HTTPError as exc:
            warnings.append(_warning("link_discovery_failed", "crawl4ai", f"{current_url}: {exc}"))
            continue

        if page["status"] != "ok":
            warnings.append(_warning("link_discovery_failed", "crawl4ai", page.get("error") or current_url))
            continue

        current_entry["title"] = page.get("title") or current_entry["title"]

        for link in page.get("links", []):
            link_url = link["url"]
            normalized_link = _normalize_url(link_url)
            domain = _domain_from_url(link_url).lower()
            bare_domain = domain[4:] if domain.startswith("www.") else domain

            if same_domain_only and not _match_domain(bare_domain, [bare_root_domain]):
                continue
            if include_patterns and not _url_matches_patterns(link_url, include_patterns):
                continue
            if exclude_patterns and _url_matches_patterns(link_url, exclude_patterns):
                continue

            if normalized_link not in discovered:
                if len(discovered) >= max_urls:
                    break
                discovered[normalized_link] = {
                    "url": link_url,
                    "normalized_url": normalized_link,
                    "domain": _domain_from_url(link_url),
                    "title": link.get("title"),
                    "link_text": link.get("text"),
                    "depth": depth + 1,
                    "discovered_from": current_url,
                    "link_type": link.get("link_type", "internal"),
                }

            if depth + 1 <= max_depth and normalized_link not in visited_pages:
                queue.append((link_url, depth + 1, current_url))

    results = list(discovered.values())
    for rank, entry in enumerate(results, start=1):
        entry["rank"] = rank

    response = {
        "url": root_url,
        "results": results,
        "meta": {
            "max_urls_requested": max_urls,
            "max_depth": max_depth,
            "urls_returned": len(results),
            "pages_visited": len(visited_pages),
            "same_domain_only": same_domain_only,
            "warnings": warnings,
            "timings_ms": {
                "total": int((time.monotonic() - started) * 1000),
            },
        },
    }
    log.info(
        "map_site url=%s returned=%d visited=%d warnings=%d",
        root_url,
        len(results),
        len(visited_pages),
        len(warnings),
    )
    return response


@mcp.tool
async def map_site(
    url: str,
    max_urls: int = 25,
    max_depth: int = 1,
    include_patterns: list[str] | None = None,
    exclude_patterns: list[str] | None = None,
    same_domain_only: bool = True,
) -> str:
    response = await _map_site_impl(
        url=url,
        max_urls=max_urls,
        max_depth=max_depth,
        include_patterns=include_patterns,
        exclude_patterns=exclude_patterns,
        same_domain_only=same_domain_only,
    )
    return _format_map_results(response)


@mcp.tool
async def crawl_site(
    url: str,
    query: str | None = None,
    max_urls: int = 10,
    max_depth: int = 1,
    include_patterns: list[str] | None = None,
    exclude_patterns: list[str] | None = None,
    same_domain_only: bool = True,
    ctx: Context | None = None,
) -> str:
    """Discover and extract a bounded set of in-scope pages from a site."""
    effective_max_urls = _validate_positive_int(
        "max_urls",
        max_urls,
        maximum=min(_MAX_MAP_URLS, _MAX_EXTRACT_URLS),
    )
    started = time.monotonic()

    mapped = await _map_site_impl(
        url=url,
        max_urls=effective_max_urls,
        max_depth=max_depth,
        include_patterns=include_patterns,
        exclude_patterns=exclude_patterns,
        same_domain_only=same_domain_only,
    )
    mapped_urls = [result["url"] for result in mapped["results"]]
    extracted = await _extract_urls_impl(urls=mapped_urls, query=query, ctx=ctx)

    extract_by_normalized_url = {
        result["normalized_url"]: result
        for result in extracted["results"]
    }

    results: list[dict] = []
    for mapped_result in mapped["results"]:
        extracted_result = extract_by_normalized_url.get(mapped_result["normalized_url"])
        if not extracted_result:
            continue
        combined = {
            **mapped_result,
            "status": extracted_result["status"],
            "content_type": extracted_result.get("content_type"),
            "content": extracted_result.get("content", ""),
            "top_chunks": extracted_result.get("top_chunks", []),
            "cached": extracted_result.get("cached", False),
            "error": extracted_result.get("error"),
        }
        if extracted_result.get("title"):
            combined["title"] = extracted_result["title"]
        results.append(combined)

    if query:
        results.sort(
            key=lambda result: max(
                (chunk.get("score") or 0.0) for chunk in result.get("top_chunks", [])
            ) if result.get("top_chunks") else float("-inf"),
            reverse=True,
        )
        for rank, result in enumerate(results, start=1):
            result["rank"] = rank

    response = {
        "url": mapped["url"],
        "query": query.strip() if query else None,
        "results": results,
        "meta": {
            "max_urls_requested": effective_max_urls,
            "max_depth": max_depth,
            "urls_discovered": mapped["meta"]["urls_returned"],
            "urls_succeeded": extracted["meta"]["urls_succeeded"],
            "urls_failed": extracted["meta"]["urls_failed"],
            "same_domain_only": same_domain_only,
            "warnings": [
                *mapped["meta"].get("warnings", []),
            ],
            "timings_ms": {
                "map": mapped["meta"]["timings_ms"]["total"],
                "extract": extracted["meta"]["timings_ms"]["total"],
                "total": int((time.monotonic() - started) * 1000),
            },
        },
    }
    log.info(
        "crawl_site url=%s discovered=%d succeeded=%d failed=%d",
        url,
        mapped["meta"]["urls_returned"],
        extracted["meta"]["urls_succeeded"],
        extracted["meta"]["urls_failed"],
    )
    return _format_crawl_results(response)


async def _site_search_impl(
    query: str,
    site: str,
    num_results: int = 10,
    scrape_top: int = MAX_SCRAPE,
    mode: str = "balanced",
    include_domains: list[str] | None = None,
    exclude_domains: list[str] | None = None,
    categories: list[str] | None = None,
    language: str | None = None,
    safesearch: int | None = None,
    ctx: Context | None = None,
) -> dict:
    """Search within a specific website or domain.

    Prepends 'site:<domain>' to the query and runs the full search pipeline.
    Useful for searching documentation sites, GitHub repos, or any specific domain.

    Args:
        query: The search query.
        site: Domain to search within (e.g. 'github.com', 'docs.python.org').
        num_results: Number of search results to fetch (default 10).
        scrape_top: Number of top results to scrape for full content (default 5).
    """
    normalized_site = site.strip()
    if not normalized_site:
        raise ValueError("site must not be empty")
    scoped_query = f"site:{normalized_site} {query}"
    return await _web_search_impl(
        query=scoped_query,
        num_results=num_results,
        scrape_top=scrape_top,
        mode=mode,
        include_domains=include_domains,
        exclude_domains=exclude_domains,
        categories=categories,
        language=language,
        safesearch=safesearch,
        ctx=ctx,
    )


@mcp.tool
async def site_search(
    query: str,
    site: str,
    num_results: int = 10,
    scrape_top: int = MAX_SCRAPE,
    mode: str = "balanced",
    include_domains: list[str] | None = None,
    exclude_domains: list[str] | None = None,
    categories: list[str] | None = None,
    language: str | None = None,
    safesearch: int | None = None,
    ctx: Context | None = None,
) -> str:
    response = await _site_search_impl(
        query=query,
        site=site,
        num_results=num_results,
        scrape_top=scrape_top,
        mode=mode,
        include_domains=include_domains,
        exclude_domains=exclude_domains,
        categories=categories,
        language=language,
        safesearch=safesearch,
        ctx=ctx,
    )
    return _format_search_results(response)


_STOPWORDS = frozenset({
    "a", "about", "above", "after", "again", "against", "all", "am", "an",
    "and", "any", "are", "aren't", "as", "at", "be", "because", "been",
    "before", "being", "below", "between", "both", "but", "by", "can",
    "can't", "could", "couldn't", "did", "didn't", "do", "does", "doesn't",
    "doing", "don't", "down", "during", "each", "few", "for", "from",
    "further", "get", "got", "had", "hadn't", "has", "hasn't", "have",
    "haven't", "having", "he", "her", "here", "hers", "herself", "him",
    "himself", "his", "how", "i", "if", "in", "into", "is", "isn't", "it",
    "its", "itself", "just", "let", "like", "ll", "me", "might", "more",
    "most", "must", "mustn't", "my", "myself", "no", "nor", "not", "now",
    "of", "off", "on", "once", "only", "or", "other", "our", "ours",
    "ourselves", "out", "over", "own", "re", "s", "same", "she", "should",
    "shouldn't", "so", "some", "such", "t", "than", "that", "the", "their",
    "theirs", "them", "themselves", "then", "there", "these", "they", "this",
    "those", "through", "to", "too", "under", "until", "up", "us", "ve",
    "very", "was", "wasn't", "we", "were", "weren't", "what", "when",
    "where", "which", "while", "who", "whom", "why", "will", "with",
    "won't", "would", "wouldn't", "you", "your", "yours", "yourself",
    "yourselves", "also", "new", "one", "two", "use", "using", "used",
    "may", "well", "even", "much", "many", "make", "made", "still",
})


def _extract_keywords(content: str, max_keywords: int = 10) -> list[str]:
    """Extract top keywords from content using term frequency, excluding stopwords."""
    words = _WORD_SPLIT.split(content.lower())
    freq: dict[str, int] = {}
    for word in words:
        if len(word) < 3 or word in _STOPWORDS or word.isdigit():
            continue
        freq[word] = freq.get(word, 0) + 1
    ranked = sorted(freq, key=freq.get, reverse=True)
    return ranked[:max_keywords]


async def _find_similar_impl(
    url: str,
    num_results: int = 10,
    ctx: Context | None = None,
) -> dict:
    """Find pages similar to the given URL by extracting keywords and searching."""
    started = time.monotonic()
    warnings: list[dict] = []
    timings_ms: dict[str, int] = {}

    scrape_cache: dict[str, str | None] = {}
    if ctx:
        scrape_cache = await _ctx_get_state(ctx, STATE_SCRAPE_CACHE) or {}

    # scrape the source page
    scrape_started = time.monotonic()
    result = await _scrape(url)
    content = result["content"]
    title = result.get("title")
    timings_ms["scrape"] = int((time.monotonic() - scrape_started) * 1000)

    if not content:
        return {
            "source_url": url,
            "source_title": None,
            "results": [],
            "meta": {
                "queries_generated": 0,
                "total_candidates": 0,
                "error": "failed to scrape source URL",
                "timings_ms": {**timings_ms, "total": int((time.monotonic() - started) * 1000)},
            },
        }

    # cache the scraped content
    scrape_cache[url] = content
    if ctx:
        await _ctx_set_state(ctx, STATE_SCRAPE_CACHE, scrape_cache)

    # generate search queries from content
    keywords = _extract_keywords(content)
    queries: list[str] = []
    if title:
        queries.append(title)
    if keywords[:5]:
        queries.append(" ".join(keywords[:5]))
    if keywords[5:10]:
        queries.append(" ".join(keywords[5:10]))
    if not queries:
        queries.append(content[:200])

    # search concurrently
    search_started = time.monotonic()
    search_tasks = [_search(q, num_results=num_results) for q in queries]
    search_results = await asyncio.gather(*search_tasks, return_exceptions=True)
    timings_ms["search"] = int((time.monotonic() - search_started) * 1000)

    # merge and dedup
    all_candidates: list[dict] = []
    source_normalized = _normalize_url(url)
    for i, sr in enumerate(search_results):
        if isinstance(sr, Exception):
            warnings.append(_warning("search_failed", "searxng", f"query {i+1}: {sr}"))
            continue
        for r in sr:
            if _normalize_url(r.get("url", "")) != source_normalized:
                all_candidates.append(r)
    all_candidates = _dedup_results(all_candidates)
    total_candidates = len(all_candidates)

    if not all_candidates:
        return {
            "source_url": url,
            "source_title": title,
            "results": [],
            "meta": {
                "queries_generated": len(queries),
                "total_candidates": 0,
                "warnings": warnings,
                "timings_ms": {**timings_ms, "total": int((time.monotonic() - started) * 1000)},
            },
        }

    # rerank against source content summary
    rerank_started = time.monotonic()
    summary = content[:1000]
    candidate_snippets = [
        f"{r.get('title', '')} {r.get('content', '')}"[:500] for r in all_candidates
    ]
    try:
        scored = _rerank_scored(summary, candidate_snippets)
    except Exception as exc:
        warnings.append(_warning("rerank_failed", "flashrank", str(exc)))
        scored = [(i, 0.0) for i in range(len(all_candidates))]
    timings_ms["rerank"] = int((time.monotonic() - rerank_started) * 1000)

    # format results
    results: list[dict] = []
    for rank, (idx, score) in enumerate(scored[:num_results], 1):
        r = all_candidates[idx]
        results.append({
            "rank": rank,
            "title": r.get("title", ""),
            "url": r.get("url", ""),
            "domain": _domain_from_url(r.get("url", "")),
            "snippet": r.get("content", ""),
            "score": score,
        })

    timings_ms["total"] = int((time.monotonic() - started) * 1000)
    return {
        "source_url": url,
        "source_title": title,
        "results": results,
        "meta": {
            "queries_generated": len(queries),
            "total_candidates": total_candidates,
            "warnings": warnings,
            "timings_ms": timings_ms,
        },
    }


@mcp.tool
async def find_similar(
    url: str,
    num_results: int = 10,
    ctx: Context | None = None,
) -> str:
    """Find web pages similar to the given URL.

    Scrapes the source page, extracts keywords, searches for related content,
    and reranks candidates by similarity to the source.
    """
    url = _validate_urls([url], maximum=1)[0]
    num_results = _validate_positive_int("num_results", num_results, maximum=MAX_RESULTS)
    response = await _find_similar_impl(url=url, num_results=num_results, ctx=ctx)
    return _format_similar_results(response)


if __name__ == "__main__":
    mcp.run(
        transport="http",
        host="0.0.0.0",
        port=8000,
    )
