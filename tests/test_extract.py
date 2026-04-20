from unittest.mock import AsyncMock, patch
from io import BytesIO

import pytest
from docx import Document as DocxDocument

from tests.conftest import FakeContext, server_module

PATCH_EXTRACT_URL_DOCUMENT = "web_search_server._extract_url_document"


@pytest.mark.asyncio
async def test_extract_urls_validates_input():
    with pytest.raises(ValueError, match="urls must not be empty"):
        await server_module.extract.fn([], ctx=None)

    with pytest.raises(ValueError, match="invalid URL"):
        await server_module.extract.fn(["notaurl"], ctx=None)


@pytest.mark.asyncio
async def test_extract_urls_returns_structured_results():
    fake_ctx = FakeContext()
    extract_mock = AsyncMock(side_effect=[
        {
            "status": "ok",
            "url": "https://example.com/page",
            "content_type": "text/html",
            "title": "Example Page",
            "content": "# Example\n\nUseful extracted content.",
            "top_chunks": [{"text": "Useful extracted content.", "score": 1.0}],
            "cached": False,
        },
        {
            "status": "ok",
            "url": "https://example.com/file.pdf",
            "content_type": "application/pdf",
            "title": "Example PDF",
            "content": "## Page 1\n\nUseful PDF content.",
            "top_chunks": [],
            "cached": True,
        },
    ])

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        payload = await server_module.extract_impl(
            urls=["https://example.com/page", "https://example.com/file.pdf"],
            query="example query",
            ctx=fake_ctx,
        )

    assert payload["query"] == "example query"
    assert payload["meta"]["urls_requested"] == 2
    assert payload["meta"]["urls_succeeded"] == 2
    assert payload["meta"]["urls_failed"] == 0
    assert payload["results"][0]["status"] == "ok"
    assert payload["results"][0]["content_type"] == "text/html"
    assert payload["results"][1]["content_type"] == "application/pdf"
    assert payload["results"][1]["cached"] is True


@pytest.mark.asyncio
async def test_extract_urls_reports_partial_failures():
    extract_mock = AsyncMock(side_effect=[
        {
            "status": "ok",
            "url": "https://example.com/page",
            "content_type": "text/html",
            "title": None,
            "content": "# Example\n\nUseful extracted content.",
            "top_chunks": [],
            "cached": False,
        },
        {
            "status": "error",
            "url": "https://example.com/missing.pdf",
            "content_type": "application/pdf",
            "title": None,
            "content": "",
            "error": "404 not found",
            "top_chunks": [],
            "cached": False,
        },
    ])

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        payload = await server_module.extract_impl(
            urls=["https://example.com/page", "https://example.com/missing.pdf"],
        )

    assert payload["meta"]["urls_succeeded"] == 1
    assert payload["meta"]["urls_failed"] == 1
    assert payload["results"][1]["status"] == "error"
    assert payload["results"][1]["error"] == "404 not found"


@pytest.mark.asyncio
async def test_extract_urls_single_url_returns_markdown():
    fake_ctx = FakeContext()
    extract_mock = AsyncMock(return_value={
        "status": "ok",
        "url": "https://example.com/file.docx",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "file_type": "docx",
        "title": "Example Doc",
        "content": "Example content",
        "top_chunks": [],
        "cached": False,
    })

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        payload = await server_module.extract.fn(
            ["https://example.com/file.docx"],
            query="example query",
            ctx=fake_ctx,
        )

    assert "example query" in payload
    assert "https://example.com/file.docx" in payload
    assert "Example Doc" in payload


def test_docx_bytes_to_markdown_extracts_paragraphs_and_tables():
    document = DocxDocument()
    document.core_properties.title = "Quarterly Report"
    document.add_paragraph("Executive summary")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Quarter"
    table.cell(0, 1).text = "Revenue"
    table.cell(1, 0).text = "Q1"
    table.cell(1, 1).text = "$1M"

    buffer = BytesIO()
    document.save(buffer)

    content, title = server_module._docx_bytes_to_markdown(buffer.getvalue())

    assert title == "Quarterly Report"
    assert "Executive summary" in content
    assert "| Quarter | Revenue |" in content


def test_pdf_extract_all_pages_returns_per_page_chunks():
    import pymupdf
    doc = pymupdf.Document()
    for _ in range(5):
        doc.new_page(width=72, height=72)
    pdf_bytes = doc.tobytes()
    # Blank pages have no extractable text
    pages, title, total_pages = server_module._pdf_extract_all_pages(pdf_bytes)
    assert isinstance(pages, list)
    assert len(pages) == 0
    assert total_pages == 5


@pytest.mark.asyncio
async def test_extract_pdf_returns_per_page_chunks_with_metadata():
    fake_pages = [
        {"page": i, "content": f"Content for page {i}"}
        for i in range(1, 11)
    ]

    fake_resp = AsyncMock()
    fake_resp.content = b"fake pdf bytes"
    fake_resp.raise_for_status = lambda: None
    fake_client = AsyncMock()
    fake_client.get = AsyncMock(return_value=fake_resp)
    fake_client.__aenter__ = AsyncMock(return_value=fake_client)
    fake_client.__aexit__ = AsyncMock(return_value=None)

    with (
        patch("web_search_server._pdf_extract_all_pages", return_value=(fake_pages, "Test PDF", 10)),
        patch("web_search_server.httpx.AsyncClient", return_value=fake_client),
    ):
        result = await server_module._extract_pdf_document("https://example.com/manual.pdf")

    assert result["title"] == "Test PDF"
    assert result["total_pages"] == 10
    assert result["pages_returned"] > 0
    assert "## Page 1" in result["content"]


@pytest.mark.asyncio
async def test_extract_pdf_reranks_pages_with_query():
    fake_pages = [
        {"page": i, "content": f"Content for page {i}"}
        for i in range(1, 6)
    ]

    fake_resp = AsyncMock()
    fake_resp.content = b"fake pdf bytes"
    fake_resp.raise_for_status = lambda: None
    fake_client = AsyncMock()
    fake_client.get = AsyncMock(return_value=fake_resp)
    fake_client.__aenter__ = AsyncMock(return_value=fake_client)
    fake_client.__aexit__ = AsyncMock(return_value=None)

    # Reranker returns pages in reverse relevance order
    def fake_rerank(_query, docs):
        return [(i, 1.0 - i * 0.1) for i in reversed(range(len(docs)))]

    with (
        patch("web_search_server._pdf_extract_all_pages", return_value=(fake_pages, "Test", 5)),
        patch("web_search_server.httpx.AsyncClient", return_value=fake_client),
        patch("web_search_server._rerank_scored", side_effect=fake_rerank),
    ):
        result = await server_module._extract_pdf_document(
            "https://example.com/manual.pdf",
            query="relevant content",
        )

    # Reranked content should have score annotations
    assert "(score:" in result["content"]
    assert result["pages_returned"] > 0


@pytest.mark.asyncio
async def test_extract_urls_surfaces_pdf_pagination_metadata():
    extract_mock = AsyncMock(return_value={
        "status": "ok",
        "url": "https://example.com/manual.pdf",
        "content_type": "application/pdf",
        "file_type": "pdf",
        "title": "Manual",
        "content": "## Page 1\n\nContent",
        "total_pages": 50,
        "pages_returned": 3,
        "top_chunks": [],
        "cached": False,
    })

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        payload = await server_module.extract_impl(
            urls=["https://example.com/manual.pdf"],
        )

    result = payload["results"][0]
    assert result["total_pages"] == 50
    assert result["pages_returned"] == 3


@pytest.mark.asyncio
async def test_extract_cache_hit_preserves_pdf_pagination_metadata():
    """Cache hits must carry total_pages / pages_returned / file_type through."""
    cache = server_module._new_cache()
    cache["https://example.com/manual.pdf"] = {
        "status": "ok",
        "url": "https://example.com/manual.pdf",
        "content_type": "application/pdf",
        "file_type": "pdf",
        "title": "Manual",
        "content": "## Page 1\n\nContent",
        "total_pages": 50,
        "pages_returned": 3,
    }

    result = await server_module._extract_url_document(
        "https://example.com/manual.pdf", query=None, cache=cache,
    )

    assert result["cached"] is True
    assert result["file_type"] == "pdf"
    assert result["total_pages"] == 50
    assert result["pages_returned"] == 3


def test_guess_file_type_supports_docx_and_text_formats():
    assert server_module._guess_file_type(
        "https://example.com/file.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) == "docx"
    assert server_module._guess_file_type("https://example.com/data.json", "application/json") == "json"
    assert server_module._guess_file_type("https://example.com/notes.txt", "text/plain; charset=utf-8") == "text"


# ---------------------------------------------------------------------------
# Truncation signal + offset continuation
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_extract_markdown_signals_truncation_with_next_offset():
    """When content is truncated, markdown footer must tell the LLM
    how to continue (offset=N).

    _extract_url_document returns content already sliced to <= MAX
    plus the full total_chars alongside — simulate that shape.
    """
    full_length = 24000
    sliced = "x" * 8000
    extract_mock = AsyncMock(return_value={
        "status": "ok",
        "url": "https://example.com/long",
        "content_type": "text/html",
        "file_type": "html",
        "title": "Long Page",
        "content": sliced,
        "total_chars": full_length,
    })

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        markdown = await server_module.extract.fn(["https://example.com/long"])

    assert "chars shown" in markdown
    assert "offset=8000" in markdown, f"truncation footer missing offset hint:\n{markdown}"


@pytest.mark.asyncio
async def test_extract_no_signal_when_content_fits():
    """Short content that wasn't truncated must not emit a truncation footer."""
    short = "tiny content"
    extract_mock = AsyncMock(return_value={
        "status": "ok",
        "url": "https://example.com/short",
        "content_type": "text/html",
        "file_type": "html",
        "title": "Short",
        "content": short,
        "total_chars": len(short),
    })

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        markdown = await server_module.extract.fn(["https://example.com/short"])

    assert "chars shown" not in markdown
    assert "end of document" not in markdown


@pytest.mark.asyncio
async def test_extract_offset_bypasses_rerank_and_slices_raw():
    """offset>0 returns the raw content slice, skipping query rerank."""
    long_content = "A" * 20000
    fake_cache = server_module._new_cache()
    fake_cache["https://example.com/long"] = {
        "status": "ok",
        "url": "https://example.com/long",
        "content_type": "text/html",
        "file_type": "html",
        "title": "Long",
        "content": long_content,
        "total_chars": len(long_content),
    }

    result = await server_module._extract_url_document(
        "https://example.com/long",
        query="something",
        cache=fake_cache,
        offset=8000,
    )
    assert result["content"] == long_content[8000:16000]
    # Offset skips rerank — no top_chunks.
    assert result["top_chunks"] == []


@pytest.mark.asyncio
async def test_extract_rejects_negative_offset():
    with pytest.raises(ValueError, match="offset must be >= 0"):
        await server_module.extract_impl(
            urls=["https://example.com/a"], offset=-1,
        )


@pytest.mark.asyncio
async def test_extract_offset_reaching_end_shows_end_marker():
    long_content = "B" * 12000
    extract_mock = AsyncMock(return_value={
        "status": "ok",
        "url": "https://example.com/page",
        "content_type": "text/html",
        "file_type": "html",
        "title": "Page",
        "content": long_content[8000:],  # simulating the slice returned
        "total_chars": 12000,
    })

    with patch(PATCH_EXTRACT_URL_DOCUMENT, extract_mock):
        markdown = await server_module.extract.fn(
            ["https://example.com/page"], offset=8000,
        )

    assert "end of document" in markdown
    assert "12,000 chars total" in markdown
